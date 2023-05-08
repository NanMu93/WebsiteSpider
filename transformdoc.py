import os
import struct

import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import random
import re
from collections import deque
import json
import pymongo

def save_article(self, parent_dir, art_time, u_name, art_content=None):
    """
    保存文章
    :param parent_dir:
    :param art_time:
    :param art_content:
    :return:返回标题作为图片名
    """
    document = Document()
    art_r_time = art_time
    art_r_time, art_y_time = self.tranform_time(art_time)
    if not os.path.exists(parent_dir + "/" + u_name + "/" + art_y_time):
        os.makedirs(parent_dir + "/" + u_name + "/" + art_y_time)
    if not os.path.exists(parent_dir + "/" + u_name + "/" + art_y_time + "/" + art_r_time):
        os.makedirs(parent_dir + "/" + u_name + "/" + art_y_time + "/" + art_r_time)
    if art_content is None:
        docxname = art_r_time + str(random.randint(0, 9999)) + "_"
    else:
        art_list = art_content.split("\n")
        if len(art_list) > 1:
            if len(art_list[0]) > 20:
                docxname = art_list[0].split("，")[0]
                self.w_docx_content(document, art_list[1:])
            else:
                docxname = art_list[0]
                self.w_docx_title(document, art_list[0])
                self.w_docx_content(document, art_list[1:])
        else:
            if len(art_content) > 20:
                docxname = art_content.split("，")[0]
            else:
                docxname = art_content
            self.w_docx_content(document, art_content.split("\n"))
        document.save(parent_dir + "/" + u_name + "/" + art_y_time + '/' + art_r_time + '/' + docxname + '.docx')
    return docxname + "_", parent_dir + "/" + u_name + "/" + art_y_time + '/' + art_r_time


class TransformDocx:
    def __init__(self, mongo_uri, mongo_db, mongo_collection, parentpath):
        self.mongo_uri = mongo_uri
        self.mongo_db = mongo_db
        self.mongo_collection = mongo_collection
        self.client = pymongo.MongoClient(self.mongo_uri)
        self.db = self.client[self.mongo_db]
        self.path = parentpath

    def tranform(self):
        """
        转换主方法
        :return:
        """
        for item in self.db[self.mongo_collection].find():
            doc = Document()
            docname = item['title']
            # 定义不允许出现在文件名中的特殊字符
            illegal_chars = r'[\\/:*?"<>|]'

            # 使用正则表达式替换非法字符为空字符串
            docname = re.sub(illegal_chars, '', docname)
            self.w_docx_title(doc, item['title'])
            content = self.remove_div_tags(item['content'])
            self.w_docx_content(doc, content.split('\n'))
            try:
                if item['cover_tag'] == 1 and len(item['images']) > 1:
                    self.w_docx_picture(doc, item['images'][1:])
                    paragraph = doc.add_paragraph()
                    paragraph.add_run('缩略图：')
                    hyperlink = self.add_hyperlink(paragraph, '../images/' + item['images'][0]['path'], item['images'][0]['path'].split('/')[1], '0000FF', False)
                elif item['cover_tag'] == 1 and len(item['images']) == 1:
                    paragraph = doc.add_paragraph()
                    paragraph.add_run('缩略图：')
                    hyperlink = self.add_hyperlink(paragraph, '../images/' + item['images'][0]['path'], item['images'][0]['path'].split('/')[1], '0000FF', False)
            except:
                print("except")
                self.w_docx_picture(doc, item['images'])
            try:
                if item['files']:
                    paragraph = doc.add_paragraph()
                    paragraph.add_run('视频：')
                    video_hyperlink = self.add_hyperlink(paragraph, '../videos/' + item['files'][0]['path'], item['files'][0]['path'].split('/')[1], '0000FF', False)
                    try:
                        paragraph1 = doc.add_paragraph()
                        paragraph1.add_run('视频封面：')
                        cover_hyperlink = self.add_hyperlink(paragraph1, '../videos/' + item['files'][1]['path'], item['files'][1]['path'].split('/')[1], '0000FF', False)
                    except IndexError:
                        paragraph1.add_run('无封面图')
            except KeyError:
                pass
            self.w_docx_content(doc, ['发表时间：' + item['publish_time']])
            try:
                self.w_docx_content(doc, ['来源：' + item['source']])
            except TypeError:
                self.w_docx_content(doc, ['来源：'])
            doc.save(self.path + '大公旅讯' + '/' + docname + '.docx')

    def w_docx_title(self, document, art_title):
        """
        写入标题
        :param document:
        :param art_title:
        :return:
        """
        run = document.add_heading("", level=0).add_run(art_title.lstrip())
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)

    def w_docx_content(self, document, art_list):
        """
        写入段落
        :param document:
        :param art_list:
        :return:
        """
        for p in art_list:
            pg = document.add_paragraph()
            # 设置内容
            pg.text = p
            # 设置字号
            pg.style.font.size = Pt(14)
            pg.paragraph_format.first_line_indent = pg.style.font.size * 2
            # 设置英文字符字体
            pg.style.font.name = '黑体'
            # 设置中文字符字体
            pg.style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def w_docx_picture(self, document, pic_list):
        page_width = document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin
        for i in pic_list:
            for p in document.paragraphs:
#                img_tag = '<img src="' + i['url'] + '">'
                img_tag = 'src="' + i['url'] + '"'
                if img_tag in p.text:
                    width = Inches(2.5)  # 设置默认宽度
                    height = Inches(2.5)
                    img = p.add_run()
                    try:
                        img.add_picture(self.path + 'images/' + i['path'], width=page_width)
                    except ZeroDivisionError:
                        img.add_picture(self.path + 'images/' + i['path'], width=page_width, height=page_width)
                    except struct.error:
                        pass
                    except docx.image.exceptions.UnexpectedEndOfFileError:
                        self.add_hyperlink(p, '../images/' + i['path'], i['path'].split('/')[1], '0000FF', False)
                    for run in p.runs:
                        if img_tag in run.text:
                            # 替换字符串为空字符串
                            run.text = run.text.replace(img_tag, '')
                        if '<img alt="" >' in run.text:
                            run.text = run.text.replace('<img alt="" >', '')
                        if '<img >' in run.text:
                            run.text = run.text.replace('<img >', '')

#                    p.text = p.text.replace(img_tag, '')

    @staticmethod
    def add_hyperlink(paragraph, url, text, color, underline):
        paragraph.paragraph_format.first_line_indent = paragraph.style.font.size * 2
        # 获取文档中的关联 ID
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # 创建超链接标签
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        # 创建文本标签
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # 创建样式标签并设置颜色和下划线
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        if color is not None:
            c = docx.oxml.shared.OxmlElement('w:color')
            c.set(docx.oxml.shared.qn('w:val'), color)
            rPr.append(c)
        if not underline:
            u = docx.oxml.shared.OxmlElement('w:u')
            u.set(docx.oxml.shared.qn('w:val'), 'none')
            rPr.append(u)

        # 将标签组合并添加文本
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # 将超链接添加到段落
        paragraph._p.append(hyperlink)

        return hyperlink

    @staticmethod
    def remove_div_tags(text):
        # 使用正则表达式删除<div>标签及其内容
        cleaned_text = re.sub(r'<div.*?>|</div>|<center.*?>|</center>|<video.*?>|<source.*?>|</source>|</video>', '', text)
        return cleaned_text

    def main(self):
        """
        入口
        :return:
        """
#        input("按回车键开始任务...")
        self.tranform()
        self.client.close()


if __name__ == '__main__':
    transformdocx = TransformDocx('mongodb://localhost:27017/', 'takungpao', 'travelnews', 'F:/takungpao/')
    transformdocx.main()
